import os
from typing import List
from langchain.schema import Document
from docx import Document as DocxDocument

def parse_pdf(file_path: str) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        print("📄 PDF Text Preview:\n", text[:800], "\n--- END OF PREVIEW ---")
        return text
    except Exception as e:
        print("❗ PDF parsing failed:", str(e))
        return ""

def parse_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        print("📄 TXT File Preview:\n", text[:800], "\n--- END OF PREVIEW ---")
        return text
    except Exception as e:
        print("❗ TXT parsing failed:", str(e))
        return ""

def parse_docx(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        print("📄 DOCX Raw Text Preview:\n", text[:800], "\n--- END OF PREVIEW ---")
        return text
    except Exception as e:
        print("❗ DOCX parsing failed:", str(e))
        return ""

def process_file(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        text = parse_pdf(file_path)
    elif ext == ".txt":
        text = parse_txt(file_path)
    elif ext == ".docx":
        text = parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    if not text.strip():
        print("⚠️ Parsed text is empty")
    else:
        print(f"✅ Successfully parsed file with {len(text.split())} words.")

    return [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]
